#!/usr/bin/env python3
"""
Enhanced Resume Generator
Creates professional resumes from JSON configuration with advanced formatting.
"""

import json
import argparse
import sys
from datetime import date
from pathlib import Path
from typing import Dict, List, Any

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
import docx


class ResumeGenerator:
    """Enhanced resume generator with configuration-driven content."""
    
    def __init__(self, config_path: str = "resume_config.json"):
        """Initialize with configuration file."""
        self.config_path = Path(config_path)
        self.config = self.load_config()
        self.doc = Document()
        self.setup_styles()
    
    def load_config(self) -> Dict[str, Any]:
        """Load and validate configuration file."""
        try:
            with open(self.config_path, 'r', encoding='utf-8') as f:
                config = json.load(f)
            return config
        except FileNotFoundError:
            raise FileNotFoundError(f"Configuration file not found: {self.config_path}")
        except json.JSONDecodeError as e:
            raise ValueError(f"Invalid JSON in configuration file: {e}")
    
    def setup_styles(self):
        """Set up document styles for consistent formatting."""
        # Add custom styles
        try:
            heading_style = self.doc.styles.add_style('CustomHeading', WD_STYLE_TYPE.PARAGRAPH)
            heading_font = heading_style.font
            heading_font.name = 'Arial'
            heading_font.size = Pt(12)
            heading_font.bold = True
            heading_style.paragraph_format.space_before = Pt(12)
            heading_style.paragraph_format.space_after = Pt(6)
        except ValueError:
            # Style already exists
            pass
        
        # Set default font
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(11)
    
    def add_hyperlink(self, paragraph, text: str, url: str):
        """Add hyperlink to paragraph with proper formatting."""
        try:
            part = paragraph.part
            r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
            hyperlink = OxmlElement('w:hyperlink')
            hyperlink.set(qn('r:id'), r_id)
            
            new_run = OxmlElement('w:r')
            rPr = OxmlElement('w:rPr')
            
            # Blue color for links
            color = OxmlElement('w:color')
            color.set(qn('w:val'), '0000FF')
            rPr.append(color)
            
            # Underline
            u = OxmlElement('w:u')
            u.set(qn('w:val'), 'single')
            rPr.append(u)
            
            new_run.append(rPr)
            new_run_text = OxmlElement('w:t')
            new_run_text.text = text
            new_run.append(new_run_text)
            hyperlink.append(new_run)
            paragraph._p.append(hyperlink)
        except Exception as e:
            # Fallback to plain text if hyperlink fails
            paragraph.add_run(text)
    
    def add_header(self):
        """Add header with name and contact information."""
        personal = self.config['personal_info']
        
        # Name as main heading
        name_para = self.doc.add_heading(personal['name'], level=0)
        name_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Professional title
        title_para = self.doc.add_paragraph(personal['title'])
        title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        title_font = title_para.runs[0].font
        title_font.size = Pt(12)
        title_font.italic = True
        
        # Contact information
        contact_para = self.doc.add_paragraph()
        contact_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        contact_info = [
            f"📧 {personal['email']}",
            f"📞 {personal['phone']}",
            f"📍 {personal['location']}"
        ]
        
        contact_para.add_run(" | ".join(contact_info))
        
        # Add links paragraph
        links_para = self.doc.add_paragraph()
        links_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        if 'github' in personal and personal['github']:
            self.add_hyperlink(links_para, "GitHub", personal['github'])
            links_para.add_run(" | ")
        
        if 'linkedin' in personal and personal['linkedin']:
            self.add_hyperlink(links_para, "LinkedIn", personal['linkedin'])
    
    def add_summary(self):
        """Add professional summary section."""
        self.doc.add_heading('SUMMARY', level=1)
        self.doc.add_paragraph(self.config['summary'])
    
    def add_technical_skills(self):
        """Add technical skills section with categorized formatting."""
        self.doc.add_heading('TECHNICAL SKILLS', level=1)
        
        skills = self.config['technical_skills']
        
        # Programming Languages
        if 'programming_languages' in skills:
            lang_text = "**Programming Languages**: "
            prog_langs = skills['programming_languages']
            
            if 'expert' in prog_langs:
                lang_text += f"Expert: {', '.join(prog_langs['expert'])}; "
            if 'proficient' in prog_langs:
                lang_text += f"Proficient: {', '.join(prog_langs['proficient'])}; "
            if 'familiar' in prog_langs:
                lang_text += f"Familiar: {', '.join(prog_langs['familiar'])}"
            
            self.doc.add_paragraph(lang_text.rstrip('; '))
        
        # Libraries and Tools
        if 'libraries_and_tools' in skills:
            tools = skills['libraries_and_tools']
            tools_text = "**Libraries & Tools**: "
            
            all_tools = []
            for category, items in tools.items():
                if isinstance(items, list):
                    all_tools.extend(items)
            
            tools_text += ", ".join(all_tools)
            self.doc.add_paragraph(tools_text)
        
        # Methodologies
        if 'methodologies' in skills:
            methods_text = f"**Methodologies**: {', '.join(skills['methodologies'])}"
            self.doc.add_paragraph(methods_text)
        
        # Languages
        if 'languages' in skills:
            lang_text = f"**Languages**: {', '.join(skills['languages'])}"
            self.doc.add_paragraph(lang_text)
    
    def add_experience(self):
        """Add work experience section."""
        self.doc.add_heading('EXPERIENCE', level=1)
        
        for exp in self.config['experience']:
            for position in exp['positions']:
                # Company and position header
                header_text = f"{exp['company']} – {position['title']}"
                header_para = self.doc.add_paragraph(header_text)
                header_run = header_para.runs[0]
                header_run.font.bold = True
                
                # Location and dates
                location_text = f"{exp['location']} | {position['period']}"
                location_para = self.doc.add_paragraph(location_text)
                location_run = location_para.runs[0]
                location_run.font.italic = True
                
                # Achievements
                for achievement in position['achievements']:
                    bullet_para = self.doc.add_paragraph(f"• {achievement}")
                    bullet_para.paragraph_format.left_indent = Inches(0.25)
    
    def add_projects(self):
        """Add projects section with links."""
        self.doc.add_heading('PROJECTS', level=1)
        
        for project in self.config['projects']:
            project_para = self.doc.add_paragraph()
            project_para.add_run(f"{project['name']} (").font.bold = True
            
            # Add links
            links = project.get('links', {})
            link_parts = []
            
            for link_name, link_url in links.items():
                if link_url:
                    link_parts.append((link_name.title(), link_url))
            
            for i, (link_name, link_url) in enumerate(link_parts):
                self.add_hyperlink(project_para, link_name, link_url)
                if i < len(link_parts) - 1:
                    project_para.add_run(" | ")
            
            project_para.add_run(")")
            project_para.add_run(f"\n• {project['description']}")
    
    def add_education(self):
        """Add education section."""
        self.doc.add_heading('EDUCATION', level=1)
        
        for edu in self.config['education']:
            edu_text = f"{edu['degree']} – {edu['institution']}"
            if 'gpa' in edu:
                edu_text += f" (GPA: {edu['gpa']})"
            
            self.doc.add_paragraph(edu_text)
    
    def add_extras(self):
        """Add extras section."""
        self.doc.add_heading('EXTRAS', level=1)
        
        for extra in self.config['extras']:
            self.doc.add_paragraph(f"• {extra}")
    
    def generate_resume(self, output_path: str = None, version: str = "standard") -> str:
        """Generate complete resume document."""
        if output_path is None:
            today = date.today().isoformat()
            name = self.config['personal_info']['name'].replace(' ', '_')
            output_path = f"Resume_{name}_{version}_{today}.docx"
        
        # Add all sections
        self.add_header()
        self.add_summary()
        self.add_technical_skills()
        self.add_experience()
        self.add_projects()
        self.add_education()
        self.add_extras()
        
        # Save document
        self.doc.save(output_path)
        return output_path
    
    def generate_targeted_resume(self, target_type: str = "quant", output_path: str = None) -> str:
        """Generate targeted resume for specific job types."""
        # Customize based on target type
        if target_type == "quant":
            # Emphasize quantitative skills and trading results
            pass
        elif target_type == "tech":
            # Emphasize programming and ML skills
            pass
        elif target_type == "academic":
            # Emphasize research and teaching
            pass
        
        return self.generate_resume(output_path, f"targeted_{target_type}")


def main():
    """Main function with command-line interface."""
    parser = argparse.ArgumentParser(description="Enhanced Resume Generator")
    parser.add_argument("--config", "-c", default="resume_config.json",
                        help="Path to configuration JSON file")
    parser.add_argument("--output", "-o", 
                        help="Output file path (default: auto-generated)")
    parser.add_argument("--version", "-v", default="standard",
                        choices=["standard", "quant", "tech", "academic"],
                        help="Resume version to generate")
    parser.add_argument("--validate", action="store_true",
                        help="Validate configuration file only")
    
    args = parser.parse_args()
    
    try:
        generator = ResumeGenerator(args.config)
        
        if args.validate:
            print(f"✅ Configuration file '{args.config}' is valid")
            return
        
        if args.version == "standard":
            output_path = generator.generate_resume(args.output)
        else:
            output_path = generator.generate_targeted_resume(args.version, args.output)
        
        print(f"✅ Resume generated successfully: {output_path}")
        
    except Exception as e:
        print(f"❌ Error: {e}", file=sys.stderr)
        sys.exit(1)


if __name__ == "__main__":
    main()